from docx import Document
# from docx.shared import Inches

from lib import get_divisions, fill_line

document = Document()

"""
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
"""

values = get_divisions()

# Storing results
results = []


# hdr_cells = table.rows[0].cells

for _ in range(20):
    data = []

    table = document.add_table(rows=4, cols=4)
    
    for _ in range(4):
        # Here we create a 4x4 matrix 

        line = fill_line(values, 4)
        
        # Save the results
        results += [ l[-1] for l in line ]

        # remove the results from the line
        line = [ l[0] for l in line ]

        data.append(line)

    for line in data:
    
        row_cells = table.add_row().cells
        row_cells[0].text = line[0]
        row_cells[1].text = line[1]
        row_cells[2].text = line[2]
        row_cells[3].text = line[3]
    
    document.add_page_break()

document.save('document.docx')

# Remove duplicates    
results = set(results)
print(results)
